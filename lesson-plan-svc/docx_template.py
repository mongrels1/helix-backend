"""
docx_template.py — Structure-locked template parsing + filling.

Core guarantee: we NEVER add, remove, reorder, relabel, or reformat any
structural node (tbl / tr / tc / p / sdt). We only:
  1. write text into runs of paragraphs that are EMPTY placeholders sitting
     directly after a known field label (the template ships these blanks), and
  2. set the selected value of existing content-control dropdowns (SDT).

`structural_signature()` fingerprints the element tree (tags + nesting, no
text). We assert it is identical before and after fill — that is the machine
check behind "the structure/content of the lesson plan cannot be altered."
"""
from __future__ import annotations
import copy, hashlib, re
from typing import Optional
from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

# ---- known field labels (normalized, lowercase, no trailing punctuation) ----
LABELS = {
    "standard/expectations": "standard",
    "wida language standard": "wida",
    "today, we are learning to": "learning_target",   # sits under "Learning Targets"
    "i know i am successful when i can": "success_criteria",  # under "Success Criteria"
    "instructional resources": "instructional_resources",
    "teacher actions": "teacher_actions",
    "student actions": "student_actions",
    "academic vocabulary": "academic_vocabulary",
    "writing strategies": "writing_strategies",
    "assessment strategies": "assessment_strategies",
}

def _norm(s: str) -> str:
    return re.sub(r"[\s]+", " ", (s or "")).strip().lower().rstrip(":").strip()

# --------------------------------------------------------------------------- #
#  Structural signature
# --------------------------------------------------------------------------- #
def structural_signature(doc: Document) -> str:
    """Fingerprint of the body's element tree: tags + order + nesting, no text,
    no run properties. Two docs with the same signature have identical structure."""
    STRUCT = {"tbl", "tr", "tc", "p", "sdt", "sectPr", "tblGrid", "gridCol"}
    parts = []
    def walk(el, depth):
        tag = el.tag.split("}")[-1]
        if tag in STRUCT:
            parts.append(f"{depth}:{tag}")
            depth += 1
        for ch in el:
            walk(ch, depth)
    walk(doc.element.body, 0)
    return hashlib.sha256("|".join(parts).encode()).hexdigest()

# --------------------------------------------------------------------------- #
#  Cell / paragraph helpers
# --------------------------------------------------------------------------- #
def _iter_cells(doc: Document):
    """Yield every table cell in the document, recursing into nested tables."""
    def rec(table):
        for row in table.rows:
            for cell in row.cells:
                yield cell
                for nt in cell.tables:
                    yield from rec(nt)
    for t in doc.tables:
        yield from rec(t)

def _is_empty_para(p: Paragraph) -> bool:
    return p.text.strip() == "" and p._p.find(".//" + qn("w:sdt")) is None

def _set_para_text(p: Paragraph, text: str, italic=False):
    """Overwrite a paragraph's text WITHOUT changing its structure: reuse/keep
    the first run, clear extra run texts. Adds runs only if the paragraph has
    none (an empty <w:p/> placeholder), which does not alter the doc's
    structural signature (runs are not structural)."""
    runs = p.runs
    if not runs:
        r = p.add_run(text)
        r.italic = italic
        return
    runs[0].text = text
    runs[0].italic = italic
    for extra in runs[1:]:
        extra.text = ""

def _fill_after_label(cell, label_key: str, value: str, max_lines=1) -> bool:
    """Find the paragraph whose normalized text == a label mapping to label_key,
    then write `value` into the following empty placeholder paragraph(s)."""
    paras = cell.paragraphs
    for i, p in enumerate(paras):
        key = LABELS.get(_norm(p.text))
        # some labels are combined in one paragraph e.g. 'Learning Targets:\nToday...'
        if key != label_key:
            continue
        # write into subsequent empty paragraphs
        written = 0
        lines = value.split("\n")
        for j in range(i + 1, len(paras)):
            if written >= max_lines or written >= len(lines):
                break
            if _is_empty_para(paras[j]):
                _set_para_text(paras[j], lines[written])
                written += 1
            elif paras[j].text.strip() and _norm(paras[j].text) not in LABELS:
                break  # hit non-empty non-label content; stop
        if written == 0:
            # no blank placeholder after label -> append text onto label para run
            _set_para_text(p, p.text.rstrip() + " " + value)
        return True
    return False

# --------------------------------------------------------------------------- #
#  Content-control (SDT) dropdowns
# --------------------------------------------------------------------------- #
def _all_sdts(doc: Document):
    return doc.element.body.findall(".//" + qn("w:sdt"))

def _sdt_alias(sdt):
    props = sdt.find(qn("w:sdtPr"))
    if props is None:
        return None
    a = props.find(qn("w:alias"))
    return a.get(qn("w:val")) if a is not None else None

def _set_sdt_display(sdt, display: str):
    """Set a dropdown SDT's shown value to `display` by rewriting the text run
    inside <w:sdtContent> (keeps the control + its item list intact)."""
    content = sdt.find(qn("w:sdtContent"))
    if content is None:
        return
    ts = content.findall(".//" + qn("w:t"))
    if ts:
        ts[0].text = display
        for extra in ts[1:]:
            extra.text = ""
    else:
        # find first run and add a <w:t>
        r = content.find(".//" + qn("w:r"))
        if r is not None:
            t = r.makeelement(qn("w:t"), {})
            t.text = display
            r.append(t)

# --------------------------------------------------------------------------- #
#  Public: parse + fill
# --------------------------------------------------------------------------- #
def parse_template(path: str) -> dict:
    """Return a lightweight map of what the template contains (for the UI /
    field-map display). Positional model: main grid columns = days; day tables
    (in body order) = the same days; each day table's rows = segments."""
    doc = Document(path)
    tables = doc.tables
    days = []
    if tables:
        hdr = tables[0].rows[0].cells
        days = [c.text.strip() for c in hdr if c.text.strip()]
    # day segment tables = every top-level table except the first (grid) and last (small group)
    seg_tables = tables[1:-1] if len(tables) >= 3 else []
    segments_per_day = max((len(t.rows) for t in seg_tables), default=0)
    lesson_dd = sum(1 for s in _all_sdts(doc) if _sdt_alias(s) == "Lesson Component")
    diff_dd = sum(1 for s in _all_sdts(doc) if _sdt_alias(s) == "Differentiation")
    return {
        "days": days or ["Monday", "Tuesday", "Wednesday", "Thursday", "Friday"],
        "num_tables": len(tables),
        "segments_per_day": segments_per_day or 3,
        "lesson_component_dropdowns": lesson_dd,
        "differentiation_dropdowns": diff_dd,
        "fields_per_day": ["standard", "wida", "learning_target", "success_criteria",
                            "instructional_resources"],
        "fields_per_segment": ["lesson_component", "teacher_actions", "student_actions",
                               "academic_vocabulary", "writing_strategies",
                               "assessment_strategies", "differentiation"],
        "has_small_group_table": len(tables) >= 3,
    }

def fill_template(in_path: str, out_path: str, plan: dict) -> dict:
    """Fill the template with `plan` and write to out_path.
    `plan` shape:
      {
        "header": {...optional header fields...},
        "days": [
           {"name":"Monday",
            "standard":..,"wida":..,"learning_target":..,"success_criteria":..,
            "instructional_resources":..,
            "segments":[
               {"lesson_component":"Engage","minutes":10,
                "teacher_actions":..,"student_actions":..,
                "academic_vocabulary":..,"writing_strategies":..,
                "assessment_strategies":..,
                "differentiation":{"ESOL":..,"Gifted":..,"SWD":..}},
               ...],
           }, ... ],
        "small_group":[{"rationale":..,"identifier":..,"skill":..,"strategy":..,"activity":..}]
      }
    """
    doc = Document(in_path)
    sig_before = structural_signature(doc)

    tables = doc.tables
    grid = tables[0]
    seg_tables = tables[1:-1]
    small_group_tbl = tables[-1] if len(tables) >= 3 else None

    days_plan = plan.get("days", [])

    # ---- main grid: columns are days ----
    # row1 = Standard + WIDA ; row2 = Learning Target + Success Criteria ; row3 = Instructional Resources
    for di, day in enumerate(days_plan):
        # find the grid column for this day by header text, fallback to index
        col = di
        for ri in range(1, len(grid.rows)):
            try:
                cell = grid.rows[ri].cells[col]
            except IndexError:
                continue
            _fill_after_label(cell, "standard", day.get("standard", ""))
            _fill_after_label(cell, "wida", day.get("wida", ""))
            _fill_after_label(cell, "learning_target", day.get("learning_target", ""))
            _fill_after_label(cell, "success_criteria", day.get("success_criteria", ""))
            _fill_after_label(cell, "instructional_resources",
                              day.get("instructional_resources", ""))

    # ---- day segment tables ----
    for di, day in enumerate(days_plan):
        if di >= len(seg_tables):
            break
        st = seg_tables[di]
        segs = day.get("segments", [])
        for si, seg_row in enumerate(st.rows):
            if si >= len(segs):
                break
            seg = segs[si]
            outer_cell = seg_row.cells[0]
            # Lesson Component dropdown lives in this segment row
            for sdt in outer_cell._tc.findall(".//" + qn("w:sdt")):
                if _sdt_alias(sdt) == "Lesson Component":
                    lc = seg.get("lesson_component", "")
                    _set_sdt_display(sdt, f"Lesson Component: {lc}" if lc else "Lesson Component")
            # nested 1x3 table: col0 teacher/student, col1 differentiation, col2 vocab/writing/assess
            for nt in outer_cell.tables:
                cells = nt.rows[0].cells
                if len(cells) >= 1:
                    _fill_after_label(cells[0], "teacher_actions", seg.get("teacher_actions", ""))
                    _fill_after_label(cells[0], "student_actions", seg.get("student_actions", ""))
                if len(cells) >= 3:
                    _fill_after_label(cells[2], "academic_vocabulary", seg.get("academic_vocabulary", ""))
                    _fill_after_label(cells[2], "writing_strategies", seg.get("writing_strategies", ""))
                    _fill_after_label(cells[2], "assessment_strategies", seg.get("assessment_strategies", ""))
                # differentiation column: 3 positional compartments (ESOL, Gifted, SWD).
                # Only fill the groups the teacher selected AND that were generated.
                if len(cells) >= 2:
                    diff = seg.get("differentiation", {}) or {}
                    _fill_differentiation(cells[1], diff)

    # ---- small flexible group table ----
    if small_group_tbl is not None:
        sg = plan.get("small_group", [])
        # rows: row0 header, then data rows
        for ri, row in enumerate(small_group_tbl.rows[1:], start=0):
            if ri >= len(sg):
                break
            rec = sg[ri]
            vals = [rec.get("rationale", ""), rec.get("identifier", ""),
                    rec.get("skill", ""), rec.get("strategy", ""), rec.get("activity", "")]
            for ci, cell in enumerate(row.cells):
                if ci < len(vals) and vals[ci]:
                    # write into first paragraph of the (blank) cell
                    if cell.paragraphs:
                        _set_para_text(cell.paragraphs[0], vals[ci])

    sig_after = structural_signature(doc)
    doc.save(out_path)
    return {
        "structure_locked": sig_before == sig_after,
        "signature_before": sig_before[:12],
        "signature_after": sig_after[:12],
    }

DIFF_ORDER = ["ESOL", "Gifted", "SWD"]

def _fill_differentiation(cell, diff: dict):
    """The differentiation cell is a flat sequence of tc children:
        P(label) · SDT(1) · P · P · SDT(2) · P · P · SDT(3) · P · P · P
    The 3 SDTs are the ESOL / Gifted / SWD compartments (positional order).
    For each compartment whose group is present in `diff`, set the dropdown's
    displayed label to the group name and write the accommodation into the first
    empty <w:p> after that SDT. Groups not in `diff` (unchecked by the teacher or
    not warranted) are left at the template default — nothing is added/removed."""
    # normalize keys -> {ESOL:..,Gifted:..,SWD:..}
    norm = {}
    for k, v in (diff or {}).items():
        for g in DIFF_ORDER:
            if k.lower() == g.lower():
                norm[g] = v
    tc = cell._tc
    children = list(tc)
    sdt_positions = [i for i, ch in enumerate(children)
                     if ch.tag == qn("w:sdt") and _sdt_alias(ch) == "Differentiation"]
    for idx, pos in enumerate(sdt_positions[:3]):
        group = DIFF_ORDER[idx] if idx < len(DIFF_ORDER) else None
        if not group or group not in norm or not norm[group]:
            continue
        sdt = children[pos]
        _set_sdt_display(sdt, group)
        # find the first empty <w:p> after this sdt (before the next sdt)
        next_pos = sdt_positions[idx + 1] if idx + 1 < len(sdt_positions) else len(children)
        for j in range(pos + 1, next_pos):
            ch = children[j]
            if ch.tag == qn("w:p"):
                para = Paragraph(ch, cell)
                if _is_empty_para(para):
                    _set_para_text(para, norm[group])
                    break
